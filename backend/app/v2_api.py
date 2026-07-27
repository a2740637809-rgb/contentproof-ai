import csv
import hashlib
import io
import re
from collections import defaultdict
from datetime import datetime, timezone

from docx import Document
from fastapi import APIRouter, Depends, File, HTTPException, Response, UploadFile
from pydantic import BaseModel, Field
from sqlalchemy import func, select
from sqlalchemy.orm import Session
import httpx
import trafilatura

from app.db import get_session
from app.v2_models import (
    AnalysisRunV2,
    ResearchBrief,
    ResearchComment,
    ResearchProject,
    ResearchSource,
    ResearchTheme,
    ReviewEvent,
    ThemeMembership,
)

router = APIRouter(prefix="/api/v2", tags=["content-intelligence-v2"])


class ProjectCreate(BaseModel):
    name: str = Field(min_length=1, max_length=160)
    goal: str = Field(min_length=1, max_length=1000)


class ManualImport(BaseModel):
    article_title: str = ""
    comments: list[str] = Field(min_length=1)


class AnalysisRequest(BaseModel):
    preferred_clusters: int | None = Field(default=None, ge=2, le=12)


class ThemeUpdate(BaseModel):
    name: str | None = Field(default=None, min_length=1, max_length=160)
    status: str | None = Field(
        default=None, pattern="^(pending_review|confirmed|rejected)$"
    )


class BriefRequest(BaseModel):
    selected_theme_ids: list[int] | None = None


class WebImport(BaseModel):
    url: str = Field(pattern=r"^https?://")


class MergeRequest(BaseModel):
    source_theme_ids: list[int] = Field(min_length=1)
    name: str = Field(min_length=1, max_length=160)


class SplitRequest(BaseModel):
    name: str = Field(min_length=1, max_length=160)
    comment_ids: list[int] = Field(min_length=1)


class MoveCommentRequest(BaseModel):
    target_theme_id: int


PHONE_RE = re.compile(r"(?<!\d)(1[3-9]\d)(\d{4})(\d{4})(?!\d)")
EMAIL_RE = re.compile(r"[\w.+-]+@[\w.-]+\.[A-Za-z]{2,}")
AD_RE = re.compile(r"(加微|微信|vx|代理|返利|兼职|扫码)", re.IGNORECASE)


def utc_now():
    return datetime.now(timezone.utc)


def project_json(item: ResearchProject):
    return {
        "id": item.id,
        "name": item.name,
        "goal": item.goal,
        "stage": item.stage,
        "created_at": item.created_at,
        "updated_at": item.updated_at,
    }


def comment_json(item: ResearchComment):
    return {
        "id": item.id,
        "project_id": item.project_id,
        "source_id": item.source_id,
        "raw_text": item.raw_text,
        "cleaned_text": item.cleaned_text,
        "status": item.status,
        "exclusion_reasons": item.exclusion_reasons,
        "pii_flags": item.pii_flags,
    }


def clean_comment(text: str):
    raw = text.strip()
    flags: list[str] = []
    cleaned = PHONE_RE.sub(lambda match: f"{match.group(1)}****{match.group(3)}", raw)
    if cleaned != raw:
        flags.append("phone")
    email_cleaned = EMAIL_RE.sub("[邮箱已隐藏]", cleaned)
    if email_cleaned != cleaned:
        flags.append("email")
    reasons: list[str] = []
    if AD_RE.search(raw):
        reasons.append("advertisement")
    if not re.search(r"[\u4e00-\u9fffA-Za-z0-9]", raw):
        reasons.append("emoji_only")
    return email_cleaned, flags, reasons


def add_comments(
    session: Session, project_id: int, source_id: int, texts: list[str]
) -> tuple[int, int]:
    created = 0
    duplicates = 0
    for raw in texts:
        raw = str(raw).strip()
        if not raw:
            continue
        digest = hashlib.sha256(raw.encode("utf-8")).hexdigest()
        exists = session.scalar(
            select(ResearchComment.id).where(
                ResearchComment.project_id == project_id,
                ResearchComment.content_hash == digest,
            )
        )
        if exists is not None:
            duplicates += 1
            continue
        cleaned, flags, reasons = clean_comment(raw)
        item = ResearchComment(
            project_id=project_id,
            source_id=source_id,
            raw_text=raw,
            cleaned_text=cleaned,
            status="excluded" if reasons else "included",
            exclusion_reasons=reasons,
            pii_flags=flags,
            content_hash=digest,
        )
        session.add(item)
        session.flush()
        created += 1
    session.commit()
    return created, duplicates


def extract_public_page(url: str):
    response = httpx.get(
        url,
        follow_redirects=True,
        timeout=20,
        headers={"User-Agent": "ContentIntelligenceLab/0.3 (+local research tool)"},
    )
    response.raise_for_status()
    metadata = trafilatura.extract(
        response.text,
        output_format="json",
        with_metadata=True,
        include_comments=True,
    )
    if not metadata:
        raise ValueError("页面没有可提取的正文")
    import json

    parsed = json.loads(metadata)
    comments = parsed.get("comments") or []
    if isinstance(comments, str):
        comments = [line.strip() for line in comments.splitlines() if line.strip()]
    warnings = []
    comments_status = "success" if comments else "unavailable"
    if not comments:
        warnings.append("页面未提供公开可访问评论，请粘贴或上传评论。")
    return {
        "title": parsed.get("title") or "",
        "body": parsed.get("text") or "",
        "author": parsed.get("author") or "",
        "published_at": parsed.get("date") or "",
        "comments": comments,
        "article_status": "success",
        "comments_status": comments_status,
        "warnings": warnings,
    }


@router.post("/projects", status_code=201)
def create_project(payload: ProjectCreate, session: Session = Depends(get_session)):
    item = ResearchProject(name=payload.name, goal=payload.goal)
    session.add(item)
    session.commit()
    session.refresh(item)
    return project_json(item)


@router.get("/projects")
def list_projects(session: Session = Depends(get_session)):
    items = session.scalars(
        select(ResearchProject).order_by(ResearchProject.updated_at.desc())
    ).all()
    return {"items": [project_json(item) for item in items]}


@router.get("/projects/{project_id}")
def get_project(project_id: int, session: Session = Depends(get_session)):
    item = session.get(ResearchProject, project_id)
    if item is None:
        raise HTTPException(404, "研究项目不存在")
    result = project_json(item)
    result["comment_count"] = session.scalar(
        select(func.count(ResearchComment.id)).where(
            ResearchComment.project_id == project_id
        )
    )
    return result


@router.post("/projects/{project_id}/imports/manual", status_code=201)
def import_manual(
    project_id: int, payload: ManualImport, session: Session = Depends(get_session)
):
    project = session.get(ResearchProject, project_id)
    if project is None:
        raise HTTPException(404, "研究项目不存在")
    source = ResearchSource(
        project_id=project_id,
        kind="manual",
        title=payload.article_title,
        article_status="provided" if payload.article_title else "not_requested",
        comments_status="success",
    )
    session.add(source)
    session.commit()
    session.refresh(source)
    created, duplicates = add_comments(session, project_id, source.id, payload.comments)
    project.stage = "imported"
    project.updated_at = utc_now()
    session.commit()
    return {
        "source_id": source.id,
        "created": created,
        "duplicates": duplicates,
        "article_status": source.article_status,
        "comments_status": source.comments_status,
        "warnings": source.warnings,
    }


@router.post("/projects/{project_id}/imports/spreadsheet", status_code=201)
async def import_spreadsheet(
    project_id: int,
    file: UploadFile = File(...),
    session: Session = Depends(get_session),
):
    if session.get(ResearchProject, project_id) is None:
        raise HTTPException(404, "研究项目不存在")
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        raise HTTPException(413, "文件不能超过10MB")
    suffix = (file.filename or "").lower().rsplit(".", 1)[-1]
    rows: list[dict] = []
    if suffix == "csv":
        decoded = content.decode("utf-8-sig")
        rows = list(csv.DictReader(io.StringIO(decoded)))
    elif suffix == "xlsx":
        from openpyxl import load_workbook

        sheet = load_workbook(io.BytesIO(content), read_only=True).active
        values = list(sheet.values)
        if values:
            headers = [str(value or "") for value in values[0]]
            rows = [dict(zip(headers, row)) for row in values[1:]]
    else:
        raise HTTPException(422, "仅支持CSV或XLSX文件")
    candidates = ("评论", "评论内容", "comment", "text", "content")
    comment_column = next(
        (name for name in candidates if rows and name in rows[0]), None
    )
    if comment_column is None:
        raise HTTPException(422, "未找到评论列，请使用“评论”或“comment”列名")
    source = ResearchSource(
        project_id=project_id,
        kind="spreadsheet",
        title=file.filename or "评论表格",
        comments_status="success",
    )
    session.add(source)
    session.commit()
    session.refresh(source)
    created, duplicates = add_comments(
        session, project_id, source.id, [str(row.get(comment_column) or "") for row in rows]
    )
    return {
        "source_id": source.id,
        "created": created,
        "duplicates": duplicates,
        "article_status": "not_requested",
        "comments_status": "success",
        "warnings": [],
    }


@router.post("/projects/{project_id}/imports/web", status_code=201)
def import_web(
    project_id: int, payload: WebImport, session: Session = Depends(get_session)
):
    project = session.get(ResearchProject, project_id)
    if project is None:
        raise HTTPException(404, "研究项目不存在")
    try:
        extracted = extract_public_page(payload.url)
    except Exception as error:
        extracted = {
            "title": "",
            "body": "",
            "author": "",
            "published_at": "",
            "comments": [],
            "article_status": "failed",
            "comments_status": "not_attempted",
            "warnings": [f"网页抓取失败：{type(error).__name__}。请改用粘贴或表格导入。"],
        }
    source = ResearchSource(
        project_id=project_id,
        kind="web",
        url=payload.url,
        title=extracted["title"],
        body=extracted["body"],
        author=extracted["author"],
        published_at=extracted["published_at"],
        article_status=extracted["article_status"],
        comments_status=extracted["comments_status"],
        warnings=extracted["warnings"],
    )
    session.add(source)
    session.commit()
    session.refresh(source)
    created, duplicates = add_comments(
        session, project_id, source.id, extracted["comments"]
    )
    if created:
        project.stage = "imported"
        session.commit()
    return {
        "source_id": source.id,
        "created": created,
        "duplicates": duplicates,
        "article_status": source.article_status,
        "comments_status": source.comments_status,
        "warnings": source.warnings,
        "article": {
            "title": source.title,
            "author": source.author,
            "published_at": source.published_at,
            "body_preview": source.body[:240],
        },
    }


@router.get("/projects/{project_id}/comments")
def list_comments(
    project_id: int,
    status: str | None = None,
    session: Session = Depends(get_session),
):
    query = select(ResearchComment).where(ResearchComment.project_id == project_id)
    if status:
        query = query.where(ResearchComment.status == status)
    items = session.scalars(query.order_by(ResearchComment.id)).all()
    return {"items": [comment_json(item) for item in items], "total": len(items)}


@router.post("/comments/{comment_id}/restore")
def restore_comment(comment_id: int, session: Session = Depends(get_session)):
    item = session.get(ResearchComment, comment_id)
    if item is None:
        raise HTTPException(404, "评论不存在")
    item.status = "included"
    item.exclusion_reasons = []
    session.commit()
    return comment_json(item)


THEMES = {
    "报名与参与": ("报名", "预约", "参加", "材料", "条件"),
    "时间与地点": ("时候", "时间", "哪里", "地点", "在哪"),
    "费用与价格": ("价格", "收费", "费用", "多少钱", "免费"),
    "安全与可信": ("安全", "风险", "真假", "可信", "来源"),
}


def theme_json(item: ResearchTheme, session: Session):
    ids = list(
        session.scalars(
            select(ThemeMembership.comment_id).where(
                ThemeMembership.theme_id == item.id
            )
        ).all()
    )
    return {
        "id": item.id,
        "name": item.name,
        "summary": item.summary,
        "status": item.status,
        "comment_ids": ids,
        "cluster_label": item.cluster_label,
    }


@router.post("/projects/{project_id}/analysis", status_code=201)
def analyze_project(
    project_id: int,
    payload: AnalysisRequest,
    session: Session = Depends(get_session),
):
    comments = session.scalars(
        select(ResearchComment).where(
            ResearchComment.project_id == project_id,
            ResearchComment.status == "included",
        )
    ).all()
    if not comments:
        raise HTTPException(409, "没有可分析的评论")
    run = AnalysisRunV2(
        project_id=project_id,
        status="running",
        step="clustering",
        started_at=utc_now(),
    )
    session.add(run)
    session.commit()
    session.refresh(run)
    groups: dict[str, list[ResearchComment]] = defaultdict(list)
    for comment in comments:
        matched = False
        for name, words in THEMES.items():
            if any(word in comment.cleaned_text for word in words):
                groups[name].append(comment)
                matched = True
                break
        if not matched:
            groups["其他待研判"].append(comment)
    themes = []
    for label, (name, members) in enumerate(groups.items()):
        theme = ResearchTheme(
            project_id=project_id,
            analysis_run_id=run.id,
            name=name,
            summary=f"由{len(members)}条语义相近评论形成，等待人工核查。",
            cluster_label=label,
        )
        session.add(theme)
        session.flush()
        session.add_all(
            [ThemeMembership(theme_id=theme.id, comment_id=item.id) for item in members]
        )
        themes.append(theme)
    run.status = "completed"
    run.step = "completed"
    run.completed_at = utc_now()
    project = session.get(ResearchProject, project_id)
    project.stage = "review"
    session.commit()
    return {
        "id": run.id,
        "status": run.status,
        "step": run.step,
        "embedding_model": run.embedding_model,
        "clustering_method": run.clustering_method,
        "themes": [theme_json(item, session) for item in themes],
    }


@router.get("/projects/{project_id}/themes")
def list_themes(project_id: int, session: Session = Depends(get_session)):
    items = session.scalars(
        select(ResearchTheme)
        .where(ResearchTheme.project_id == project_id)
        .order_by(ResearchTheme.id)
    ).all()
    return {"items": [theme_json(item, session) for item in items]}


@router.patch("/themes/{theme_id}")
def update_theme(
    theme_id: int, payload: ThemeUpdate, session: Session = Depends(get_session)
):
    item = session.get(ResearchTheme, theme_id)
    if item is None:
        raise HTTPException(404, "主题不存在")
    before = {"name": item.name, "status": item.status}
    if payload.name is not None:
        item.name = payload.name
    if payload.status is not None:
        item.status = payload.status
    after = {"name": item.name, "status": item.status}
    session.add(
        ReviewEvent(
            project_id=item.project_id,
            theme_id=item.id,
            action="update",
            before=before,
            after=after,
        )
    )
    session.commit()
    return theme_json(item, session)


@router.post("/themes/{theme_id}/merge")
def merge_themes(
    theme_id: int, payload: MergeRequest, session: Session = Depends(get_session)
):
    target = session.get(ResearchTheme, theme_id)
    if target is None:
        raise HTTPException(404, "目标主题不存在")
    sources = session.scalars(
        select(ResearchTheme).where(ResearchTheme.id.in_(payload.source_theme_ids))
    ).all()
    if len(sources) != len(set(payload.source_theme_ids)):
        raise HTTPException(404, "存在无效的来源主题")
    if any(item.project_id != target.project_id for item in sources):
        raise HTTPException(409, "不能合并不同项目的主题")
    before = theme_json(target, session)
    existing = set(before["comment_ids"])
    for source in sources:
        memberships = session.scalars(
            select(ThemeMembership).where(ThemeMembership.theme_id == source.id)
        ).all()
        for membership in memberships:
            if membership.comment_id not in existing:
                membership.theme_id = target.id
                existing.add(membership.comment_id)
            else:
                session.delete(membership)
        source.status = "rejected"
    target.name = payload.name
    session.flush()
    after = theme_json(target, session)
    session.add(
        ReviewEvent(
            project_id=target.project_id,
            theme_id=target.id,
            action="merge",
            before=before,
            after=after,
        )
    )
    session.commit()
    return theme_json(target, session)


@router.post("/themes/{theme_id}/split", status_code=201)
def split_theme(
    theme_id: int, payload: SplitRequest, session: Session = Depends(get_session)
):
    source = session.get(ResearchTheme, theme_id)
    if source is None:
        raise HTTPException(404, "来源主题不存在")
    memberships = session.scalars(
        select(ThemeMembership).where(
            ThemeMembership.theme_id == source.id,
            ThemeMembership.comment_id.in_(payload.comment_ids),
        )
    ).all()
    if len(memberships) != len(set(payload.comment_ids)):
        raise HTTPException(409, "只能拆分当前主题中的评论")
    new_theme = ResearchTheme(
        project_id=source.project_id,
        analysis_run_id=source.analysis_run_id,
        name=payload.name,
        summary="由人工从原主题中拆分，等待确认。",
        status="pending_review",
        cluster_label=-2,
    )
    session.add(new_theme)
    session.flush()
    for membership in memberships:
        membership.theme_id = new_theme.id
    session.flush()
    after = theme_json(new_theme, session)
    session.add(
        ReviewEvent(
            project_id=source.project_id,
            theme_id=new_theme.id,
            action="split",
            before={"source_theme_id": source.id},
            after=after,
        )
    )
    session.commit()
    return theme_json(new_theme, session)


@router.post("/comments/{comment_id}/move")
def move_comment(
    comment_id: int,
    payload: MoveCommentRequest,
    session: Session = Depends(get_session),
):
    target = session.get(ResearchTheme, payload.target_theme_id)
    comment = session.get(ResearchComment, comment_id)
    if target is None or comment is None:
        raise HTTPException(404, "评论或目标主题不存在")
    if target.project_id != comment.project_id:
        raise HTTPException(409, "不能跨项目移动评论")
    memberships = session.scalars(
        select(ThemeMembership)
        .join(ResearchTheme, ResearchTheme.id == ThemeMembership.theme_id)
        .where(
            ThemeMembership.comment_id == comment_id,
            ResearchTheme.project_id == target.project_id,
        )
    ).all()
    before = {"theme_ids": [item.theme_id for item in memberships]}
    for membership in memberships:
        session.delete(membership)
    session.add(ThemeMembership(theme_id=target.id, comment_id=comment.id))
    session.flush()
    after = theme_json(target, session)
    session.add(
        ReviewEvent(
            project_id=target.project_id,
            theme_id=target.id,
            action="move_comment",
            before=before,
            after={"target_theme_id": target.id, "comment_id": comment.id},
        )
    )
    session.commit()
    return after


@router.get("/projects/{project_id}/review-events")
def list_review_events(project_id: int, session: Session = Depends(get_session)):
    items = session.scalars(
        select(ReviewEvent)
        .where(ReviewEvent.project_id == project_id)
        .order_by(ReviewEvent.id.desc())
    ).all()
    return {
        "items": [
            {
                "id": item.id,
                "theme_id": item.theme_id,
                "action": item.action,
                "before": item.before,
                "after": item.after,
                "created_at": item.created_at,
            }
            for item in items
        ]
    }


def brief_json(item: ResearchBrief):
    return {
        "id": item.id,
        "project_id": item.project_id,
        "version": item.version,
        "title": item.title,
        "audience": item.audience,
        "problem": item.problem,
        "angle": item.angle,
        "outline": item.outline,
        "risks": item.risks,
        "evidence_comment_ids": item.evidence_comment_ids,
        "theme_ids": item.theme_ids,
        "generation_mode": item.generation_mode,
    }


@router.post("/projects/{project_id}/briefs", status_code=201)
def create_brief(
    project_id: int, payload: BriefRequest, session: Session = Depends(get_session)
):
    query = select(ResearchTheme).where(
        ResearchTheme.project_id == project_id,
        ResearchTheme.status == "confirmed",
    )
    if payload.selected_theme_ids:
        query = query.where(ResearchTheme.id.in_(payload.selected_theme_ids))
    themes = session.scalars(query.order_by(ResearchTheme.id)).all()
    if not themes:
        raise HTTPException(409, "请先确认至少一个读者主题")
    evidence_ids = list(
        session.scalars(
            select(ThemeMembership.comment_id).where(
                ThemeMembership.theme_id.in_([item.id for item in themes])
            )
        ).all()
    )
    version = (
        session.scalar(
            select(func.max(ResearchBrief.version)).where(
                ResearchBrief.project_id == project_id
            )
        )
        or 0
    ) + 1
    theme_names = "、".join(item.name for item in themes)
    item = ResearchBrief(
        project_id=project_id,
        version=version,
        title=f"围绕“{theme_names}”的下一篇解释型报道",
        audience="提出相关疑问、需要明确行动信息的读者",
        problem=f"现有内容尚未充分回答：{theme_names}",
        angle="以读者原始问题为线索，用可核查信息逐项回应。",
        outline=[
            "读者最集中的问题是什么",
            "现有报道遗漏了哪些行动信息",
            "用事实和步骤逐项回答",
            "补充风险边界与权威来源",
        ],
        risks=["发布前核验时间、地点和参与条件", "评论只能证明读者疑问，不能替代事实来源"],
        evidence_comment_ids=sorted(set(evidence_ids)),
        theme_ids=[theme.id for theme in themes],
    )
    session.add(item)
    session.commit()
    session.refresh(item)
    return brief_json(item)


@router.get("/projects/{project_id}/briefs")
def list_briefs(project_id: int, session: Session = Depends(get_session)):
    items = session.scalars(
        select(ResearchBrief)
        .where(ResearchBrief.project_id == project_id)
        .order_by(ResearchBrief.version.desc())
    ).all()
    return {"items": [brief_json(item) for item in items]}


def brief_markdown(item: ResearchBrief, session: Session):
    comments = session.scalars(
        select(ResearchComment).where(
            ResearchComment.id.in_(item.evidence_comment_ids)
        )
    ).all()
    evidence = "\n".join(f"- [{comment.id}] {comment.raw_text}" for comment in comments)
    outline = "\n".join(f"{index}. {text}" for index, text in enumerate(item.outline, 1))
    risks = "\n".join(f"- {text}" for text in item.risks)
    return (
        f"# {item.title}\n\n## 目标读者\n{item.audience}\n\n"
        f"## 核心问题\n{item.problem}\n\n## 内容角度\n{item.angle}\n\n"
        f"## 文章结构\n{outline}\n\n## 评论证据\n{evidence}\n\n## 风险提示\n{risks}\n"
    )


@router.get("/briefs/{brief_id}/export/markdown")
def export_markdown(brief_id: int, session: Session = Depends(get_session)):
    item = session.get(ResearchBrief, brief_id)
    if item is None:
        raise HTTPException(404, "Brief不存在")
    return Response(
        brief_markdown(item, session),
        media_type="text/markdown; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="brief-v{item.version}.md"'},
    )


@router.get("/briefs/{brief_id}/export/docx")
def export_docx(brief_id: int, session: Session = Depends(get_session)):
    item = session.get(ResearchBrief, brief_id)
    if item is None:
        raise HTTPException(404, "Brief不存在")
    document = Document()
    document.add_heading(item.title, level=0)
    for heading, content in (
        ("目标读者", item.audience),
        ("核心问题", item.problem),
        ("内容角度", item.angle),
    ):
        document.add_heading(heading, level=1)
        document.add_paragraph(content)
    document.add_heading("文章结构", level=1)
    for text in item.outline:
        document.add_paragraph(text, style="List Number")
    document.add_heading("评论证据", level=1)
    comments = session.scalars(
        select(ResearchComment).where(
            ResearchComment.id.in_(item.evidence_comment_ids)
        )
    ).all()
    for comment in comments:
        document.add_paragraph(f"[{comment.id}] {comment.raw_text}", style="List Bullet")
    stream = io.BytesIO()
    document.save(stream)
    return Response(
        stream.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="brief-v{item.version}.docx"'},
    )
